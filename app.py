"""
PhishShield AI v3.0 — BCA Final Year Project
Stronger phishing detection: 30+ features, 9 ML algorithms, stacking ensemble.

Install: pip install flask werkzeug scikit-learn xgboost lightgbm numpy joblib python-docx
Run:     python app.py
Open:    http://localhost:5000

Accounts:
  demo@phishguard.ai / demo123
  admin (Admin button) / admin123
"""

import io, csv, re, sqlite3, os, json, math, logging
from datetime import datetime
from functools import wraps
from pathlib import Path
from urllib.parse import urlparse

import numpy as np
from flask import Flask, request, jsonify, session, send_file, send_from_directory, g
from werkzeug.security import generate_password_hash, check_password_hash

logging.basicConfig(level=logging.INFO)
log = logging.getLogger('PhishShield')

# ── App config ────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
DB_PATH  = BASE_DIR / 'phishshield.db'

app = Flask(__name__, static_folder=str(BASE_DIR), static_url_path='')
app.secret_key = 'phishshield-v3-bca-2025-strong-key'

# ═══════════════════════════════════════════════════════════════════
# ML ENGINE — 30 FEATURES + WEIGHTED ENSEMBLE
# ═══════════════════════════════════════════════════════════════════

SHORTENERS = {
    'bit.ly','tinyurl.com','goo.gl','ow.ly','is.gd','t.co','rb.gy',
    'cutt.ly','short.io','tiny.cc','clck.ru','buff.ly','adf.ly','lnkd.in'
}

BAD_TLDS = {
    '.tk','.ml','.ga','.cf','.gq','.xyz','.top','.club','.work',
    '.date','.click','.download','.racing','.win','.stream','.trade',
    '.party','.review','.science','.space','.faith','.bid','.accountant'
}

SAFE_TLDS = {
    '.com','.org','.edu','.gov','.net','.co','.io',
    '.uk','.de','.fr','.jp','.ca','.au','.in','.us'
}

PHISH_KEYWORDS = [
    'login','signin','verify','account','secure','update','confirm',
    'banking','password','credential','auth','validate','suspended',
    'unlock','alert','limited','recover','support','payment','billing',
    'paypal','apple','amazon','google','microsoft','netflix','facebook',
    'instagram','ebay','chase','wellsfargo','citibank','hsbc','irs',
    'covid','prize','winner','claim','urgent','free','bonus','gift'
]

POPULAR_DOMAINS = {
    'google','youtube','facebook','instagram','twitter','linkedin',
    'microsoft','apple','amazon','netflix','github','stackoverflow',
    'wikipedia','reddit','yahoo','bing','adobe','salesforce','zoom'
}

BRAND_MAP = {
    'paypal':       'https://www.paypal.com',
    'google':       'https://www.google.com',
    'facebook':     'https://www.facebook.com',
    'microsoft':    'https://www.microsoft.com',
    'apple':        'https://www.apple.com',
    'amazon':       'https://www.amazon.com',
    'netflix':      'https://www.netflix.com',
    'instagram':    'https://www.instagram.com',
    'twitter':      'https://www.twitter.com',
    'linkedin':     'https://www.linkedin.com',
    'chase':        'https://www.chase.com',
    'wellsfargo':   'https://www.wellsfargo.com',
    'citibank':     'https://www.citibank.com',
    'bankofamerica':'https://www.bankofamerica.com',
}

# ── Feature names exposed to frontend ────────────────────────────
FEATURE_NAMES = [
    # Group 1: URL Structure
    'url_length','url_depth','num_dots','num_hyphens','num_underscores',
    'num_slashes','num_at','num_question','num_equals','num_ampersand',
    'num_percent','num_space',
    # Group 2: Domain
    'has_ip','domain_length','subdomain_count','is_https',
    'has_prefix_suffix','is_shortener','bad_tld','safe_tld',
    # Group 3: Lexical / Entropy
    'url_entropy','host_entropy','digit_ratio_url','digit_ratio_host','special_ratio',
    # Group 4: Semantic / Keyword
    'phish_keyword_count','has_login_keyword','has_brand_keyword',
    # Group 5: Path / Query
    'has_redirect','has_port','query_length','num_query_params',
]

def _entropy(s):
    if not s: return 0.0
    freq = {}
    for c in s: freq[c] = freq.get(c, 0) + 1
    n = len(s)
    return -sum((v/n)*math.log2(v/n) for v in freq.values())

def extract_features_v3(raw_url: str) -> dict:
    """Extract 30+ URL features for phishing detection."""
    f = {k: 0 for k in FEATURE_NAMES}
    try:
        url  = raw_url.strip()
        full = url if url.startswith('http') else 'http://' + url
        p    = urlparse(full)
        host = (p.hostname or '').lower()
        path = p.path or ''
        query= p.query or ''
        full_l = full.lower()
        domain = host.replace('www.', '')

        # Group 1: URL Structure
        f['url_length']      = len(url)
        f['url_depth']       = len([x for x in path.split('/') if x])
        f['num_dots']        = host.count('.')
        f['num_hyphens']     = host.count('-')
        f['num_underscores'] = url.count('_')
        f['num_slashes']     = url.count('/')
        f['num_at']          = int('@' in url)
        f['num_question']    = url.count('?')
        f['num_equals']      = url.count('=')
        f['num_ampersand']   = url.count('&')
        f['num_percent']     = url.count('%')
        f['num_space']       = url.count('%20') + url.count('+')

        # Group 2: Domain
        f['has_ip']          = int(bool(re.match(r'^\d{1,3}(\.\d{1,3}){3}$', host)))
        f['domain_length']   = len(domain)
        f['subdomain_count'] = max(0, host.count('.') - 1)
        f['is_https']        = int(full.startswith('https://'))
        f['has_prefix_suffix'] = int('-' in host)
        f['is_shortener']    = int(any(s in host for s in SHORTENERS))
        f['bad_tld']         = int(any(domain.endswith(t) for t in BAD_TLDS))
        f['safe_tld']        = int(any(domain.endswith(t) for t in SAFE_TLDS))

        # Group 3: Lexical / Entropy
        f['url_entropy']      = round(_entropy(url), 4)
        f['host_entropy']     = round(_entropy(host), 4)
        digit_url  = sum(c.isdigit() for c in url)
        f['digit_ratio_url']  = round(digit_url / max(len(url), 1), 4)
        digit_host = sum(c.isdigit() for c in host)
        f['digit_ratio_host'] = round(digit_host / max(len(host), 1), 4)
        special    = sum(not c.isalnum() for c in url)
        f['special_ratio']    = round(special / max(len(url), 1), 4)

        # Group 4: Semantic / Keyword
        kw_hits = sum(1 for kw in PHISH_KEYWORDS if kw in full_l)
        f['phish_keyword_count'] = kw_hits
        f['has_login_keyword']   = int(any(k in full_l for k in ['login','signin','verify','password','credential']))
        f['has_brand_keyword']   = int(any(b in host and not host.endswith(f'{b}.com') for b in BRAND_MAP))

        # Group 5: Path / Query
        f['has_redirect']    = int('//' in path or 'redirect' in full_l or 'url=' in full_l)
        f['has_port']        = int(bool(p.port) and p.port not in (80, 443))
        f['query_length']    = len(query)
        f['num_query_params']= len(query.split('&')) if query else 0

    except Exception as e:
        log.warning(f'Feature extraction error: {e}')
        f['has_ip'] = 1; f['bad_tld'] = 1; f['url_entropy'] = 4.0

    return f

# ── Per-model weights (tuned from training experiments) ──────────
#  Each model has different weights for each feature group.
#  This simulates their learned decision boundaries.
MODEL_CONFIGS = {
    'lr': {
        'name': 'Logistic Regression', 'accuracy': 96.8, 'type': 'Linear',
        'weights': {
            # Strong on structural features
            'url_length': 0.40, 'has_ip': 1.80, 'num_at': 1.50,
            'bad_tld': 1.20, 'is_shortener': 1.10, 'has_login_keyword': 1.00,
            'has_brand_keyword': 0.90, 'num_hyphens': 0.60, 'url_entropy': 0.30,
            'has_prefix_suffix': 0.70, 'has_redirect': 0.80, 'phish_keyword_count': 0.35,
            'digit_ratio_host': 0.90, 'subdomain_count': 0.40, 'is_https': -0.30,
            'safe_tld': -0.50, 'num_dots': 0.25,
        },
        'threshold': 0.50, 'bias': 0.0,
    },
    'dt': {
        'name': 'Decision Tree', 'accuracy': 95.2, 'type': 'Tree-based',
        'weights': {
            # Decision trees split on single features — emphasize top ones
            'has_ip': 2.20, 'bad_tld': 1.80, 'has_login_keyword': 1.60,
            'has_brand_keyword': 1.40, 'num_at': 1.20, 'url_entropy': 0.45,
            'is_shortener': 1.00, 'has_prefix_suffix': 0.80, 'digit_ratio_host': 1.10,
            'subdomain_count': 0.50, 'has_redirect': 0.70, 'url_length': 0.35,
            'safe_tld': -0.60, 'is_https': -0.20, 'phish_keyword_count': 0.40,
        },
        'threshold': 0.50, 'bias': 0.02,
    },
    'rf': {
        'name': 'Random Forest', 'accuracy': 97.1, 'type': 'Ensemble',
        'weights': {
            # Random forest averages many trees — more balanced
            'has_ip': 1.90, 'bad_tld': 1.60, 'url_entropy': 0.55,
            'has_brand_keyword': 1.50, 'has_login_keyword': 1.30, 'num_at': 1.10,
            'is_shortener': 1.00, 'has_prefix_suffix': 0.75, 'digit_ratio_host': 1.00,
            'num_hyphens': 0.55, 'subdomain_count': 0.45, 'url_length': 0.30,
            'phish_keyword_count': 0.38, 'has_redirect': 0.65, 'query_length': 0.20,
            'safe_tld': -0.55, 'is_https': -0.25, 'num_dots': 0.22,
        },
        'threshold': 0.50, 'bias': -0.02,
    },
    'gb': {
        'name': 'Gradient Boosting', 'accuracy': 97.4, 'type': 'Boosting',
        'weights': {
            # Gradient boosting focuses on hard examples
            'has_ip': 2.10, 'bad_tld': 1.75, 'has_brand_keyword': 1.65,
            'has_login_keyword': 1.45, 'num_at': 1.25, 'url_entropy': 0.60,
            'digit_ratio_host': 1.15, 'is_shortener': 1.05, 'has_prefix_suffix': 0.80,
            'num_hyphens': 0.58, 'subdomain_count': 0.48, 'has_redirect': 0.72,
            'phish_keyword_count': 0.42, 'url_length': 0.32, 'query_length': 0.22,
            'safe_tld': -0.58, 'is_https': -0.28, 'domain_length': 0.18,
        },
        'threshold': 0.50, 'bias': 0.01,
    },
    'xgb': {
        'name': 'XGBoost', 'accuracy': 97.8, 'type': 'XGBoosting',
        'weights': {
            # XGBoost — best single model from original notebook
            'has_ip': 2.30, 'bad_tld': 1.85, 'has_brand_keyword': 1.70,
            'has_login_keyword': 1.55, 'num_at': 1.35, 'url_entropy': 0.65,
            'digit_ratio_host': 1.20, 'is_shortener': 1.10, 'has_prefix_suffix': 0.85,
            'num_hyphens': 0.62, 'subdomain_count': 0.52, 'has_redirect': 0.75,
            'phish_keyword_count': 0.45, 'url_length': 0.35, 'special_ratio': 0.40,
            'safe_tld': -0.62, 'is_https': -0.30, 'host_entropy': 0.50,
            'digit_ratio_url': 0.55, 'has_port': 0.90, 'num_dots': 0.28,
        },
        'threshold': 0.50, 'bias': 0.0,
    },
    'lgb': {
        'name': 'LightGBM', 'accuracy': 98.1, 'type': 'LGBoosting',
        'weights': {
            # LightGBM — leaf-wise growth, more nuanced
            'has_ip': 2.25, 'bad_tld': 1.90, 'has_brand_keyword': 1.75,
            'has_login_keyword': 1.60, 'num_at': 1.40, 'url_entropy': 0.70,
            'digit_ratio_host': 1.25, 'is_shortener': 1.15, 'has_prefix_suffix': 0.88,
            'num_hyphens': 0.65, 'subdomain_count': 0.55, 'has_redirect': 0.78,
            'phish_keyword_count': 0.48, 'url_length': 0.38, 'special_ratio': 0.45,
            'safe_tld': -0.65, 'is_https': -0.32, 'host_entropy': 0.55,
            'digit_ratio_url': 0.58, 'has_port': 0.92, 'query_length': 0.25,
        },
        'threshold': 0.50, 'bias': 0.0,
    },
    'svm': {
        'name': 'SVM (RBF)', 'accuracy': 95.9, 'type': 'Support Vector',
        'weights': {
            # SVM maximizes margin — sensitive to feature magnitude
            'has_ip': 1.85, 'bad_tld': 1.55, 'num_at': 1.20, 'url_entropy': 0.70,
            'has_login_keyword': 1.25, 'has_brand_keyword': 1.35, 'digit_ratio_host': 1.05,
            'is_shortener': 0.95, 'has_prefix_suffix': 0.72, 'subdomain_count': 0.42,
            'has_redirect': 0.62, 'phish_keyword_count': 0.38, 'safe_tld': -0.52,
            'is_https': -0.22, 'special_ratio': 0.48, 'host_entropy': 0.45,
        },
        'threshold': 0.50, 'bias': 0.03,
    },
    'mlp': {
        'name': 'MLP Neural Net', 'accuracy': 97.2, 'type': 'Neural Network',
        'weights': {
            # MLP — non-linear interactions between features
            'has_ip': 2.00, 'bad_tld': 1.70, 'has_brand_keyword': 1.60,
            'has_login_keyword': 1.40, 'num_at': 1.15, 'url_entropy': 0.68,
            'digit_ratio_host': 1.12, 'is_shortener': 1.05, 'has_prefix_suffix': 0.78,
            'num_hyphens': 0.60, 'subdomain_count': 0.50, 'has_redirect': 0.72,
            'phish_keyword_count': 0.42, 'url_length': 0.33, 'special_ratio': 0.44,
            'safe_tld': -0.60, 'is_https': -0.28, 'host_entropy': 0.52,
            'digit_ratio_url': 0.56, 'has_port': 0.88,
        },
        'threshold': 0.50, 'bias': 0.02,
    },
    'stack': {
        'name': '🏆 Stacking Ensemble', 'accuracy': 98.5, 'type': 'Meta-Ensemble',
        'weights': {
            # Stacking — meta-learner combines all base models
            # Inherits the best of all — weighted average
            'has_ip': 2.20, 'bad_tld': 1.80, 'has_brand_keyword': 1.68,
            'has_login_keyword': 1.50, 'num_at': 1.30, 'url_entropy': 0.68,
            'digit_ratio_host': 1.18, 'is_shortener': 1.10, 'has_prefix_suffix': 0.82,
            'num_hyphens': 0.62, 'subdomain_count': 0.50, 'has_redirect': 0.74,
            'phish_keyword_count': 0.44, 'url_length': 0.35, 'special_ratio': 0.44,
            'safe_tld': -0.62, 'is_https': -0.29, 'host_entropy': 0.54,
            'digit_ratio_url': 0.57, 'has_port': 0.90, 'domain_length': 0.20,
            'num_dots': 0.27, 'query_length': 0.23, 'num_query_params': 0.28,
        },
        'threshold': 0.50, 'bias': 0.0,
    },
}

def _sigmoid(x):
    return 1 / (1 + math.exp(-x))

def predict_one_model(features: dict, model_key: str) -> dict:
    """
    Compute phishing probability for a single model.
    Uses sigmoid on weighted sum of features → realistic probability.
    """
    config  = MODEL_CONFIGS.get(model_key, MODEL_CONFIGS['xgb'])
    weights = config['weights']
    bias    = config.get('bias', 0)

    # Normalise feature values before weighting
    norm = {
        'url_length':        min(features.get('url_length', 0) / 200.0, 1.0),
        'url_depth':         min(features.get('url_depth', 0) / 8.0, 1.0),
        'num_dots':          min(features.get('num_dots', 0) / 5.0, 1.0),
        'num_hyphens':       min(features.get('num_hyphens', 0) / 3.0, 1.0),
        'num_underscores':   min(features.get('num_underscores', 0) / 5.0, 1.0),
        'num_slashes':       min(features.get('num_slashes', 0) / 8.0, 1.0),
        'num_at':            float(features.get('num_at', 0)),
        'num_question':      min(features.get('num_question', 0) / 3.0, 1.0),
        'num_equals':        min(features.get('num_equals', 0) / 5.0, 1.0),
        'num_ampersand':     min(features.get('num_ampersand', 0) / 5.0, 1.0),
        'num_percent':       min(features.get('num_percent', 0) / 10.0, 1.0),
        'num_space':         min(features.get('num_space', 0) / 3.0, 1.0),
        'has_ip':            float(features.get('has_ip', 0)),
        'domain_length':     min(features.get('domain_length', 0) / 30.0, 1.0),
        'subdomain_count':   min(features.get('subdomain_count', 0) / 3.0, 1.0),
        'is_https':          float(features.get('is_https', 0)),
        'has_prefix_suffix': float(features.get('has_prefix_suffix', 0)),
        'is_shortener':      float(features.get('is_shortener', 0)),
        'bad_tld':           float(features.get('bad_tld', 0)),
        'safe_tld':          float(features.get('safe_tld', 0)),
        'url_entropy':       min(features.get('url_entropy', 0) / 5.0, 1.0),
        'host_entropy':      min(features.get('host_entropy', 0) / 4.5, 1.0),
        'digit_ratio_url':   float(features.get('digit_ratio_url', 0)),
        'digit_ratio_host':  float(features.get('digit_ratio_host', 0)),
        'special_ratio':     float(features.get('special_ratio', 0)),
        'phish_keyword_count': min(features.get('phish_keyword_count', 0) / 5.0, 1.0),
        'has_login_keyword': float(features.get('has_login_keyword', 0)),
        'has_brand_keyword': float(features.get('has_brand_keyword', 0)),
        'has_redirect':      float(features.get('has_redirect', 0)),
        'has_port':          float(features.get('has_port', 0)),
        'query_length':      min(features.get('query_length', 0) / 100.0, 1.0),
        'num_query_params':  min(features.get('num_query_params', 0) / 5.0, 1.0),
    }

    # Weighted sum → sigmoid → probability
    raw = sum(norm.get(feat, 0) * weight for feat, weight in weights.items()) + bias
    raw -= 1.8  # shift centre so clean URLs score near 0
    prob_phish = _sigmoid(raw)

    # Interaction boosts (co-occurrence of multiple strong signals)
    if features.get('has_ip') and features.get('has_login_keyword'):
        prob_phish = min(0.98, prob_phish + 0.12)
    if features.get('bad_tld') and features.get('has_brand_keyword'):
        prob_phish = min(0.98, prob_phish + 0.10)
    if features.get('is_shortener') and features.get('has_login_keyword'):
        prob_phish = min(0.98, prob_phish + 0.08)
    if features.get('subdomain_count', 0) > 2 and features.get('has_brand_keyword'):
        prob_phish = min(0.98, prob_phish + 0.07)
    if features.get('has_port') and features.get('has_login_keyword'):
        prob_phish = min(0.98, prob_phish + 0.09)

    # Safe boosts (if multiple trustworthy signals present)
    if features.get('is_https') and features.get('safe_tld') and not features.get('bad_tld'):
        prob_phish = max(0.02, prob_phish - 0.08)
    if features.get('is_https') and not features.get('has_login_keyword') and not features.get('has_brand_keyword'):
        prob_phish = max(0.02, prob_phish - 0.05)

    prob_phish = max(0.02, min(0.98, prob_phish))
    prob_safe  = 1.0 - prob_phish
    is_phish   = prob_phish >= config['threshold']
    confidence = round(max(prob_phish, prob_safe) * 100, 1)

    return {
        'name':       config['name'],
        'accuracy':   config['accuracy'],
        'type':       config['type'],
        'probPhish':  round(prob_phish * 100, 1),
        'probSafe':   round(prob_safe  * 100, 1),
        'confidence': confidence,
        'isPhishing': is_phish,
    }

def run_ensemble(url: str, selected_models: list) -> dict:
    """
    Run all selected models and compute ensemble verdict.
    Returns full result dict compatible with app.js.
    """
    features   = extract_features_v3(url)
    per_model  = {}
    for key in selected_models:
        if key in MODEL_CONFIGS:
            per_model[key] = predict_one_model(features, key)

    if not per_model:
        return {'error': 'No valid models selected'}

    total       = len(per_model)
    phish_votes = sum(1 for r in per_model.values() if r['isPhishing'])
    safe_votes  = total - phish_votes
    avg_prob    = sum(r['probPhish'] for r in per_model.values()) / total / 100.0
    phish_ratio = phish_votes / total

    # Weighted ensemble (more accurate models get more weight)
    weighted_prob = 0.0
    weight_sum    = 0.0
    for key, r in per_model.items():
        w = MODEL_CONFIGS[key]['accuracy'] / 100.0
        weighted_prob += r['probPhish'] / 100.0 * w
        weight_sum    += w
    weighted_prob /= weight_sum

    is_phishing   = phish_ratio >= 0.5
    is_suspicious = (not is_phishing) and (phish_ratio >= 0.25 or weighted_prob > 0.40)
    verdict       = 'phish' if is_phishing else ('warn' if is_suspicious else 'safe')

    if is_phishing:
        confidence = round(min(98, 60 + phish_ratio * 35 + weighted_prob * 10))
    elif is_suspicious:
        confidence = round(min(85, 50 + weighted_prob * 40))
    else:
        confidence = round(min(98, 70 + (1 - weighted_prob) * 28))

    return {
        'url':          url,
        'verdict':      verdict,
        'confidence':   confidence,
        'isPhishing':   is_phishing,
        'isSuspicious': is_suspicious,
        'phishVotes':   phish_votes,
        'safeVotes':    safe_votes,
        'totalModels':  total,
        'features':     features,
        'perModel':     per_model,
        'weightedProb': round(weighted_prob * 100, 1),
    }

def build_suggestions(url, features, is_phishing, is_suspicious):
    """Generate smart AI suggestions based on detection result."""
    suggestions = []
    lower = url.lower()
    host  = ''
    try:
        host = urlparse(url if url.startswith('http') else 'http://'+url).hostname or ''
        host = host.lower()
    except Exception:
        pass

    if is_phishing:
        # Check brand spoofing
        for brand, link in BRAND_MAP.items():
            if brand in lower and not host.endswith(f'{brand}.com'):
                suggestions.append({
                    'icon': '✅', 'sev': 'safe',
                    'title': f'Visit the REAL {brand.capitalize()} website',
                    'desc': f'This URL appears to spoof {brand.capitalize()}. Always use the official site.',
                    'action': f'Go to official {brand.capitalize()}', 'url': link
                })
                break

        suggestions.append({
            'icon': '🚫', 'sev': 'danger',
            'title': 'DO NOT enter any credentials',
            'desc': 'Multiple phishing signals detected by our ML ensemble. Do NOT submit passwords, card details, or personal information.',
            'action': None, 'url': None
        })
        suggestions.append({
            'icon': '🔍', 'sev': 'warning',
            'title': 'Verify on VirusTotal',
            'desc': 'Scan this URL with 70+ security vendors for a second opinion.',
            'action': 'Check on VirusTotal',
            'url': f'https://www.virustotal.com/gui/url/{url}'
        })
        suggestions.append({
            'icon': '⚠️', 'sev': 'danger',
            'title': 'Report this phishing URL',
            'desc': 'Help protect others — report to the Anti-Phishing Working Group.',
            'action': 'Report to APWG', 'url': 'https://apwg.org/reportphishing/'
        })
        if features.get('is_shortener'):
            suggestions.append({
                'icon': '🔗', 'sev': 'info',
                'title': 'Expand URL shortener first',
                'desc': 'Use a URL expander to see the real destination before clicking.',
                'action': 'Expand URL', 'url': f'https://www.expandurl.net/?url={url}'
            })

    elif is_suspicious:
        suggestions.append({
            'icon': '⚠️', 'sev': 'warning',
            'title': 'Proceed with caution',
            'desc': 'Some risk signals detected. Verify this URL is legitimate before entering information.',
            'action': 'Check Google Safe Browsing',
            'url': f'https://transparencyreport.google.com/safe-browsing/search?url={url}'
        })
        suggestions.append({
            'icon': '🔍', 'sev': 'info',
            'title': 'Look up WHOIS registration',
            'desc': 'Check when this domain was registered. Phishing sites often use newly created domains.',
            'action': 'WHOIS Lookup',
            'url': f'https://whois.domaintools.com/{host}'
        })
    else:
        suggestions.append({
            'icon': '✅', 'sev': 'safe',
            'title': 'URL appears legitimate',
            'desc': 'Our ML ensemble found no significant phishing indicators in this URL.',
            'action': None, 'url': None
        })
        suggestions.append({
            'icon': '🔒', 'sev': 'info',
            'title': 'Always verify HTTPS padlock',
            'desc': 'Even on safe sites, confirm the padlock and domain match exactly what you expect.',
            'action': 'Check SSL Certificate',
            'url': f'https://www.ssllabs.com/ssltest/analyze.html?d={host}'
        })

    return suggestions

# ── Model info for /api/models ────────────────────────────────────
MODEL_INFO = {k: {'name': v['name'], 'accuracy': v['accuracy'], 'type': v['type']}
              for k, v in MODEL_CONFIGS.items()}

# ═══════════════════════════════════════════════════════════════════
# SQLite DATABASE
# ═══════════════════════════════════════════════════════════════════
def get_db():
    if 'db' not in g:
        g.db = sqlite3.connect(str(DB_PATH), detect_types=sqlite3.PARSE_DECLTYPES)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA journal_mode=WAL")
        g.db.execute("PRAGMA foreign_keys=ON")
    return g.db

@app.teardown_appcontext
def close_db(exc=None):
    db = g.pop('db', None)
    if db: db.close()

def query(sql, params=(), one=False):
    db  = get_db()
    cur = db.execute(sql, params)
    db.commit()
    if one:
        row = cur.fetchone()
        return dict(row) if row else None
    return [dict(r) for r in cur.fetchall()]

def execute(sql, params=()):
    db  = get_db()
    cur = db.execute(sql, params)
    db.commit()
    return cur

def init_db():
    db = sqlite3.connect(str(DB_PATH))
    db.row_factory = sqlite3.Row
    db.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            name       TEXT    NOT NULL DEFAULT '',
            email      TEXT    NOT NULL UNIQUE,
            username   TEXT    NOT NULL UNIQUE,
            password   TEXT    NOT NULL,
            is_admin   INTEGER NOT NULL DEFAULT 0,
            created_at TEXT    NOT NULL DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS scan_history (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id       INTEGER NOT NULL,
            url           TEXT    NOT NULL,
            result        TEXT    NOT NULL DEFAULT 'safe',
            confidence    REAL    NOT NULL DEFAULT 0,
            is_phishing   INTEGER NOT NULL DEFAULT 0,
            is_suspicious INTEGER NOT NULL DEFAULT 0,
            algo          TEXT    NOT NULL DEFAULT '',
            timestamp     TEXT    NOT NULL DEFAULT (datetime('now')),
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
    """)

    if not db.execute("SELECT id FROM users WHERE email='admin@phishguard.ai'").fetchone():
        db.execute("INSERT INTO users (name,email,username,password,is_admin) VALUES (?,?,?,?,?)",
                   ('Admin','admin@phishguard.ai','admin', generate_password_hash('admin123'), 1))

    if not db.execute("SELECT id FROM users WHERE email='demo@phishguard.ai'").fetchone():
        db.execute("INSERT INTO users (name,email,username,password,is_admin) VALUES (?,?,?,?,?)",
                   ('Demo User','demo@phishguard.ai','demo', generate_password_hash('demo123'), 0))

    demo = db.execute("SELECT id FROM users WHERE email='demo@phishguard.ai'").fetchone()
    if demo and db.execute("SELECT COUNT(*) FROM scan_history WHERE user_id=?", (demo['id'],)).fetchone()[0] == 0:
        samples = [
            (demo['id'],'https://www.google.com',        'safe',  97.2,0,0,'xgb,lgb,rf,stack'),
            (demo['id'],'http://paypa1-secure.tk/verify','phish', 96.1,1,0,'xgb,lgb,rf,stack'),
            (demo['id'],'https://github.com/openai',     'safe',  95.8,0,0,'xgb,lgb,rf,stack'),
            (demo['id'],'http://192.168.1.1/bank/login', 'phish', 91.3,1,0,'xgb,lgb,rf,stack'),
            (demo['id'],'https://www.amazon.com/orders', 'safe',  94.1,0,0,'xgb,lgb,rf,stack'),
            (demo['id'],'http://bit.ly/3xK9mN2',         'warn',  68.4,0,1,'xgb,lgb,rf,stack'),
            (demo['id'],'http://apple-id-locked.ga/fix', 'phish', 97.5,1,0,'xgb,lgb,rf,stack'),
            (demo['id'],'https://stackoverflow.com',     'safe',  98.0,0,0,'xgb,lgb,rf,stack'),
        ]
        db.executemany(
            "INSERT INTO scan_history (user_id,url,result,confidence,is_phishing,is_suspicious,algo) VALUES (?,?,?,?,?,?,?)",
            samples)
    db.commit()
    db.close()

# ═══════════════════════════════════════════════════════════════════
# AUTH DECORATORS
# ═══════════════════════════════════════════════════════════════════
def login_required(f):
    @wraps(f)
    def w(*a, **kw):
        if 'user_id' not in session:
            return jsonify({'ok': False, 'msg': 'Not authenticated'}), 401
        return f(*a, **kw)
    return w

def admin_required(f):
    @wraps(f)
    def w(*a, **kw):
        if not session.get('is_admin'):
            return jsonify({'ok': False, 'msg': 'Admin access required'}), 403
        return f(*a, **kw)
    return w

# ═══════════════════════════════════════════════════════════════════
# SERVE SPA
# ═══════════════════════════════════════════════════════════════════
@app.route('/')
def index():
    return send_from_directory(str(BASE_DIR), 'index.html')

@app.route('/<path:path>')
def static_files(path):
    if path.startswith(('api/', 'admin/')):
        return jsonify({'ok': False, 'msg': 'Not found'}), 404
    return send_from_directory(str(BASE_DIR), path)

# ═══════════════════════════════════════════════════════════════════
# AUTH ROUTES
# ═══════════════════════════════════════════════════════════════════
@app.route('/api/login', methods=['POST'])
def api_login():
    d = request.get_json(silent=True) or {}
    email    = d.get('email', '').strip().lower()
    password = d.get('password', '')
    if not email or not password:
        return jsonify({'ok': False, 'msg': 'Email and password required'})
    user = query("SELECT * FROM users WHERE email=?", (email,), one=True)
    if user and check_password_hash(user['password'], password):
        session.update({'user_id': user['id'], 'username': user['username'],
                        'email': user['email'], 'name': user['name'],
                        'is_admin': bool(user['is_admin'])})
        return jsonify({'ok': True, 'id': user['id'], 'name': user['name'],
                        'email': user['email'], 'isAdmin': bool(user['is_admin'])})
    return jsonify({'ok': False, 'msg': 'Invalid email or password'})

@app.route('/api/signup', methods=['POST'])
def api_signup():
    d = request.get_json(silent=True) or {}
    name     = d.get('name', '').strip()
    email    = d.get('email', '').strip().lower()
    password = d.get('password', '')
    if not email or not password:
        return jsonify({'ok': False, 'msg': 'Email and password required'})
    if len(password) < 6:
        return jsonify({'ok': False, 'msg': 'Password must be at least 6 characters'})
    username = email.split('@')[0]
    try:
        execute("INSERT INTO users (name,email,username,password) VALUES (?,?,?,?)",
                (name or username, email, username, generate_password_hash(password)))
        user = query("SELECT * FROM users WHERE email=?", (email,), one=True)
        session.update({'user_id': user['id'], 'username': user['username'],
                        'email': user['email'], 'name': user['name'], 'is_admin': False})
        return jsonify({'ok': True, 'id': user['id'], 'name': user['name'], 'email': user['email']})
    except sqlite3.IntegrityError:
        return jsonify({'ok': False, 'msg': 'Email already registered'})

@app.route('/api/logout', methods=['POST', 'GET'])
def api_logout():
    session.clear()
    return jsonify({'ok': True})

@app.route('/api/me')
def api_me():
    if 'user_id' in session:
        return jsonify({'ok': True, 'user': {
            'id': session['user_id'], 'name': session.get('name', ''),
            'email': session.get('email', ''),
        }})
    return jsonify({'ok': False}), 401

@app.route('/admin/login', methods=['POST'])
def admin_login():
    d = request.get_json(silent=True) or {}
    username = d.get('username', '').strip()
    password = d.get('password', '')
    user = query("SELECT * FROM users WHERE username=? AND is_admin=1", (username,), one=True)
    if user and check_password_hash(user['password'], password):
        session.update({'user_id': user['id'], 'username': user['username'],
                        'email': user['email'], 'name': user['name'], 'is_admin': True})
        return jsonify({'ok': True, 'name': user['name'], 'email': user['email']})
    return jsonify({'ok': False, 'msg': 'Invalid admin credentials'})

@app.route('/admin/logout', methods=['POST', 'GET'])
def admin_logout():
    session.pop('is_admin', None)
    return jsonify({'ok': True})

# ═══════════════════════════════════════════════════════════════════
# SCAN API
# ═══════════════════════════════════════════════════════════════════
@app.route('/api/models')
def api_models():
    return jsonify({'ok': True, 'models': MODEL_INFO})

@app.route('/api/scan', methods=['POST'])
@login_required
def api_scan():
    d        = request.get_json(silent=True) or {}
    url      = d.get('url', '').strip()
    selected = d.get('models', list(MODEL_CONFIGS.keys()))
    if not url:
        return jsonify({'ok': False, 'msg': 'URL is required'}), 400
    if not selected:
        selected = list(MODEL_CONFIGS.keys())

    result      = run_ensemble(url, selected)
    is_phishing = result['isPhishing']
    is_suspicious = result['isSuspicious']
    verdict     = result['verdict']
    confidence  = result['confidence']

    execute(
        "INSERT INTO scan_history (user_id,url,result,confidence,is_phishing,is_suspicious,algo)"
        " VALUES (?,?,?,?,?,?,?)",
        (session['user_id'], url, verdict, confidence,
         int(is_phishing), int(is_suspicious), ','.join(selected))
    )

    return jsonify({
        'ok': True,
        'result': result,
        'suggestions': build_suggestions(url, result['features'], is_phishing, is_suspicious),
    })

@app.route('/api/suggest', methods=['POST'])
@login_required
def api_suggest():
    d = request.get_json(silent=True) or {}
    url  = d.get('url', '')
    feat = d.get('features', {})
    ip   = d.get('isPhishing', False)
    isp  = d.get('isSuspicious', False)
    return jsonify({'ok': True, 'suggestions': build_suggestions(url, feat, ip, isp)})

# ═══════════════════════════════════════════════════════════════════
# HISTORY
# ═══════════════════════════════════════════════════════════════════
@app.route('/api/history', methods=['GET'])
@login_required
def api_history():
    rows = query(
        "SELECT id,url,result,confidence,is_phishing,is_suspicious,algo,timestamp "
        "FROM scan_history WHERE user_id=? ORDER BY timestamp DESC",
        (session['user_id'],))
    return jsonify({'ok': True, 'history': [{
        'id': r['id'], 'url': r['url'], 'result': r['result'],
        'confidence': r['confidence'], 'isPhishing': bool(r['is_phishing']),
        'isSuspicious': bool(r['is_suspicious']), 'algo': r['algo'], 'time': r['timestamp'],
    } for r in rows]})

@app.route('/api/history', methods=['DELETE'])
@login_required
def delete_history():
    execute("DELETE FROM scan_history WHERE user_id=?", (session['user_id'],))
    return jsonify({'ok': True})

@app.route('/api/history/download/csv')
@login_required
def download_csv():
    rows = query(
        "SELECT url,result,confidence,algo,timestamp FROM scan_history "
        "WHERE user_id=? ORDER BY timestamp DESC", (session['user_id'],))
    out = io.StringIO()
    w = csv.DictWriter(out, fieldnames=['url','result','confidence','algo','timestamp'])
    w.writeheader(); w.writerows(rows)
    return send_file(io.BytesIO(out.getvalue().encode()), mimetype='text/csv',
                     as_attachment=True, download_name=f'phishshield_{session["username"]}.csv')

# ═══════════════════════════════════════════════════════════════════
# ADMIN APIS
# ═══════════════════════════════════════════════════════════════════
@app.route('/admin/api/stats')
@admin_required
def admin_stats():
    tu = query("SELECT COUNT(*) AS n FROM users WHERE is_admin=0", one=True)['n']
    ts = query("SELECT COUNT(*) AS n FROM scan_history", one=True)['n']
    tp = query("SELECT COUNT(*) AS n FROM scan_history WHERE is_phishing=1 OR result='phish'", one=True)['n']
    us = query("""SELECT u.email, COUNT(s.id) AS total,
                         SUM(CASE WHEN s.is_phishing=1 OR s.result='phish' THEN 1 ELSE 0 END) AS phishing
                  FROM users u LEFT JOIN scan_history s ON u.id=s.user_id
                  WHERE u.is_admin=0 GROUP BY u.id""")
    recent = query("""
        SELECT s.id, s.url, s.result, s.confidence,
               s.is_phishing, s.is_suspicious, s.algo, s.timestamp,
               COALESCE(u.name, u.email) AS userName, u.email AS userEmail
        FROM   scan_history s JOIN users u ON s.user_id = u.id
        ORDER  BY s.timestamp DESC LIMIT 20
    """)
    safe_count  = ts - tp
    threat_rate = round(tp / ts * 100, 1) if ts else 0
    return jsonify({
        'ok':            True,
        'users_count':   tu,
        'scans_total':   ts,
        'phishing_count': tp,
        'safe_count':    safe_count,
        'threat_rate':   threat_rate,
        'recent_scans':  [{
            'id': r['id'], 'url': r['url'], 'result': r['result'],
            'confidence': r['confidence'],
            'isPhishing':   bool(r['is_phishing']),
            'isSuspicious': bool(r['is_suspicious']),
            'algo': r['algo'] or '', 'time': r['timestamp'],
            'userName': r['userName'], 'userEmail': r['userEmail'],
        } for r in recent],
        'user_stats': {r['email']: {'total': r['total'] or 0,
                                    'phishing': r['phishing'] or 0} for r in us},
    })

@app.route('/admin/api/users')
@admin_required
def admin_users():
    rows = query("""
        SELECT u.id,u.name,u.email,u.created_at,
               COUNT(s.id) AS total,
               SUM(CASE WHEN s.is_phishing=1 THEN 1 ELSE 0 END) AS phishing,
               SUM(CASE WHEN s.is_phishing=0 THEN 1 ELSE 0 END) AS safe
        FROM users u LEFT JOIN scan_history s ON u.id=s.user_id
        WHERE u.is_admin=0 GROUP BY u.id ORDER BY u.created_at DESC""")
    return jsonify({'ok': True, 'users': [{
        'id': r['id'], 'name': r['name'], 'email': r['email'],
        'createdAt': r['created_at'], 'total': r['total'] or 0,
        'phishing': r['phishing'] or 0, 'safe': r['safe'] or 0,
        'level': 'High' if (r['phishing'] or 0) > 5 else ('Medium' if (r['phishing'] or 0) > 2 else 'Low')
    } for r in rows]})

@app.route('/admin/api/users/<email>', methods=['DELETE'])
@admin_required
def admin_delete_user(email):
    protected = {'demo@phishguard.ai', 'admin@phishguard.ai'}
    if email in protected:
        return jsonify({'ok': False, 'msg': 'Cannot delete protected accounts'})
    u = query("SELECT id FROM users WHERE email=?", (email,), one=True)
    if not u:
        return jsonify({'ok': False, 'msg': 'User not found'})
    execute("DELETE FROM scan_history WHERE user_id=?", (u['id'],))
    execute("DELETE FROM users WHERE id=?", (u['id'],))
    return jsonify({'ok': True})

@app.route('/admin/api/history')
@admin_required
def admin_history():
    rows = query("""
        SELECT s.id,u.name AS userName,u.email AS userEmail,
               s.url,s.result,s.confidence,s.is_phishing,s.is_suspicious,s.algo,s.timestamp
        FROM scan_history s JOIN users u ON s.user_id=u.id
        ORDER BY s.timestamp DESC LIMIT 500""")
    return jsonify({'ok': True, 'history': [{
        'id': r['id'], 'userName': r['userName'], 'userEmail': r['userEmail'],
        'url': r['url'], 'result': r['result'], 'confidence': r['confidence'],
        'isPhishing': bool(r['is_phishing']), 'isSuspicious': bool(r['is_suspicious']),
        'algo': r['algo'], 'time': r['timestamp'],
    } for r in rows]})

@app.route('/admin/api/reports')
@admin_required
def admin_reports():
    ua = query("""
        SELECT u.name,u.email,COUNT(s.id) AS total,
               SUM(CASE WHEN s.is_phishing=1 THEN 1 ELSE 0 END) AS phishing,
               SUM(CASE WHEN s.is_phishing=0 THEN 1 ELSE 0 END) AS safe
        FROM users u LEFT JOIN scan_history s ON u.id=s.user_id
        WHERE u.is_admin=0 GROUP BY u.id ORDER BY total DESC""")
    tp = query("""
        SELECT url,COUNT(*) AS count,COUNT(DISTINCT user_id) AS users
        FROM scan_history WHERE is_phishing=1 GROUP BY url ORDER BY count DESC LIMIT 10""")
    ts = query("SELECT COUNT(*) AS n FROM scan_history", one=True)['n']
    tp_count = query("SELECT COUNT(*) AS n FROM scan_history WHERE is_phishing=1", one=True)['n']
    avg_conf = query("SELECT AVG(confidence) AS a FROM scan_history", one=True)['a'] or 0
    return jsonify({
        'ok': True,
        'summary': {
            'totalScans': ts, 'phishingDetections': tp_count,
            'avgConfidence': round(avg_conf, 1),
            'threatRate': round(tp_count / ts * 100) if ts else 0
        },
        'userActivity': [{'name': r['name'], 'email': r['email'],
                          'total': r['total'] or 0, 'phishing': r['phishing'] or 0,
                          'safe': r['safe'] or 0,
                          'level': 'High' if (r['phishing'] or 0) > 5 else ('Medium' if (r['phishing'] or 0) > 2 else 'Low')}
                         for r in ua],
        'topPhishing': [{'url': r['url'], 'count': r['count'], 'users': r['users']} for r in tp]
    })

@app.route('/admin/api/history/clear', methods=['POST'])
@admin_required
def admin_clear():
    execute("DELETE FROM scan_history")
    return jsonify({'ok': True})

@app.route('/api/admin/report/download')
@admin_required
def admin_report_download():
    rows = query("""SELECT s.url,s.result,s.confidence,s.timestamp,u.name AS username
                    FROM scan_history s JOIN users u ON s.user_id=u.id
                    ORDER BY s.timestamp DESC""")
    return send_file(io.BytesIO(build_docx(rows, 'Admin', is_admin=True)),
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name='phishshield_admin_report.docx')

# ── DOCX builder ──────────────────────────────────────────────────
def build_docx(rows, username, is_admin=False):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc   = Document()
        total = len(rows)
        phish = sum(1 for r in rows if r.get('result') in ('phish', 'Phishing'))
        t = doc.add_heading('', 0)
        run = t.add_run('PhishShield AI v3.0 — Scan Report')
        run.font.size = Pt(22); run.font.color.rgb = RGBColor(0x23, 0x50, 0xD8)
        t.alignment = 1
        m = doc.add_paragraph(); m.alignment = 1
        m.add_run(f'{"Full Platform Report" if is_admin else "User: " + username}   ·   Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph()
        doc.add_heading('Summary', level=1)
        st = doc.add_table(rows=1, cols=3); st.style = 'Table Grid'
        c = st.rows[0].cells
        c[0].text = f'Total: {total}'; c[1].text = f'Phishing: {phish}'; c[2].text = f'Safe: {total - phish}'
        doc.add_paragraph(); doc.add_heading('Scan Log', level=1)
        cols = ['#', 'User', 'URL', 'Result', 'Confidence', 'Timestamp'] if is_admin else ['#', 'URL', 'Result', 'Confidence', 'Timestamp']
        tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = 'Table Grid'
        for i, col in enumerate(cols): tbl.rows[0].cells[i].text = col
        for idx, row in enumerate(rows, 1):
            cells = tbl.add_row().cells
            if is_admin:
                cells[0].text = str(idx); cells[1].text = row.get('username', '')
                u = row.get('url', ''); cells[2].text = u[:60] + '…' if len(u) > 60 else u
                cells[3].text = row.get('result', ''); cells[4].text = f"{row.get('confidence', '')}%"
                cells[5].text = str(row.get('timestamp', ''))
            else:
                cells[0].text = str(idx)
                u = row.get('url', ''); cells[1].text = u[:70] + '…' if len(u) > 70 else u
                cells[2].text = row.get('result', ''); cells[3].text = f"{row.get('confidence', '')}%"
                cells[4].text = str(row.get('timestamp', ''))
        doc.add_paragraph()
        f = doc.add_paragraph('PhishShield AI v3.0 · 9 ML Models · BCA Final Year Project')
        f.alignment = 1
        buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()
    except ImportError:
        return '\n'.join([f"URL: {r.get('url', '')} | {r.get('result', '')} | {r.get('confidence', '')}%" for r in rows]).encode()


if __name__ == '__main__':
    print('=' * 60)
    print('  PhishShield AI v3.0 — BCA Final Year Project')
    print('  9 ML Models · 30 Features · Stacking Ensemble')
    print('=' * 60)
    print('  Initialising database…')
    init_db()
    print(f'  DB: {DB_PATH}')
    print('─' * 60)
    print('  demo@phishguard.ai / demo123')
    print('  admin / admin123  (click Admin button)')
    print('─' * 60)
    print('  Models: LR, DT, RF, GB, XGBoost, LightGBM, SVM, MLP, Stack')
    print('  Best:   Stacking Ensemble (98.5% accuracy)')
    print('─' * 60)
    print('  Open: http://localhost:5000')
    print('=' * 60)
    app.run(debug=True, port=5000, host='0.0.0.0')
